"""
IC Memo Generator (Bonus) - Generates a Word document Investment Committee memo
from a completed due diligence workflow result.
"""
import asyncio
from datetime import datetime

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from exercises.agentic_due_diligence import run_due_diligence
from app.services.integration.portfolio_data_service import portfolio_data_service


def _fmt_dt(value) -> str:
    """Format a datetime or None to a readable string."""
    if value is None:
        return "N/A"
    if isinstance(value, datetime):
        return value.strftime("%Y-%m-%d %H:%M:%S UTC")
    return str(value)


def _fmt_score(value, decimals: int = 1) -> str:
    if value is None:
        return "N/A"
    try:
        return f"{float(value):.{decimals}f}"
    except (TypeError, ValueError):
        return str(value)


async def generate_ic_memo(company_id: str, assessment_type: str = "full") -> str:
    # Look up company name and ticker for display
    company_name = company_id
    company_ticker = company_id
    try:
        views = await portfolio_data_service.get_portfolio_view("growth_fund_v")
        for v in views:
            if v.company_id == company_id:
                company_name = v.name
                company_ticker = v.ticker
                break
    except Exception:
        pass

    result = await run_due_diligence(company_id, assessment_type)

    doc = Document()

    # ── Title ──────────────────────────────────────────────────────────────────
    title = doc.add_heading(f"Investment Committee Memo — {company_name} ({company_ticker})", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Assessment metadata ────────────────────────────────────────────────────
    doc.add_heading("Assessment Overview", level=1)
    meta_table = doc.add_table(rows=3, cols=2)
    meta_table.style = "Light List Accent 1"
    rows = meta_table.rows
    rows[0].cells[0].text = "Assessment Type"
    rows[0].cells[1].text = result.get("assessment_type", assessment_type)
    rows[1].cells[0].text = "Started At"
    rows[1].cells[1].text = _fmt_dt(result.get("started_at"))
    rows[2].cells[0].text = "Completed At"
    rows[2].cells[1].text = _fmt_dt(result.get("completed_at"))

    # ── Scores ─────────────────────────────────────────────────────────────────
    doc.add_heading("Org-AI-R Scores", level=1)
    scoring = result.get("scoring_result") or {}
    score_table = doc.add_table(rows=4, cols=2)
    score_table.style = "Light List Accent 1"
    rows = score_table.rows
    rows[0].cells[0].text = "Org-AI-R Score"
    rows[0].cells[1].text = _fmt_score(scoring.get("org_air"))
    rows[1].cells[0].text = "VR Score"
    rows[1].cells[1].text = _fmt_score(scoring.get("vr_score"))
    rows[2].cells[0].text = "HR Score"
    rows[2].cells[1].text = _fmt_score(scoring.get("hr_score"))
    rows[3].cells[0].text = "Synergy Score"
    rows[3].cells[1].text = _fmt_score(scoring.get("synergy_score"))

    # ── HITL approval ──────────────────────────────────────────────────────────
    doc.add_heading("HITL Approval", level=1)
    hitl_table = doc.add_table(rows=3, cols=2)
    hitl_table.style = "Light List Accent 1"
    rows = hitl_table.rows
    rows[0].cells[0].text = "Requires Approval"
    rows[0].cells[1].text = str(result.get("requires_approval", False))
    rows[1].cells[0].text = "Approval Status"
    rows[1].cells[1].text = result.get("approval_status") or "N/A"
    rows[2].cells[0].text = "Approved By"
    rows[2].cells[1].text = result.get("approved_by") or "N/A"

    # ── Evidence justifications ────────────────────────────────────────────────
    doc.add_heading("Evidence Justifications", level=1)
    evidence = result.get("evidence_justifications") or {}
    justifications = evidence.get("justifications") or {}
    if justifications:
        for dimension, data in justifications.items():
            doc.add_heading(dimension.replace("_", " ").title(), level=2)
            dim_table = doc.add_table(rows=2, cols=2)
            dim_table.style = "Light List Accent 1"
            rows = dim_table.rows
            rows[0].cells[0].text = "Score"
            rows[0].cells[1].text = _fmt_score(data.get("score") if isinstance(data, dict) else None)
            rows[1].cells[0].text = "Evidence Strength"
            rows[1].cells[1].text = (
                str(data.get("strength", "N/A")) if isinstance(data, dict) else "N/A"
            )
            rubric = data.get("rubric", "") if isinstance(data, dict) else ""
            if rubric:
                doc.add_paragraph(str(rubric))
    else:
        doc.add_paragraph("No evidence justifications available.")

    # ── Value creation / gap analysis ─────────────────────────────────────────
    doc.add_heading("Value Creation Plan", level=1)
    vc_plan = result.get("value_creation_plan") or {}
    gap_data = vc_plan.get("gap_analysis") or {}
    if gap_data:
        scalar_keys = ["target_org_air", "gap_count", "top_priority", "estimated_investment", "projected_ebitda_pct"]
        for key in scalar_keys:
            if key in gap_data:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(f"{key.replace('_', ' ').title()}: ").bold = True
                p.add_run(str(gap_data[key]))
        gaps = gap_data.get("gaps") or []
        if gaps:
            doc.add_heading("Gap Analysis by Dimension", level=2)
            for gap in gaps:
                if isinstance(gap, dict):
                    dim = gap.get("dimension", "").replace("_", " ").title()
                    p = doc.add_paragraph(style="List Bullet")
                    p.add_run(f"{dim}: ").bold = True
                    p.add_run(
                        f"Current {gap.get('current_score', 'N/A')} → Target {gap.get('target_score', 'N/A')} "
                        f"| Gap: {gap.get('gap', 'N/A')} pts | Priority: {gap.get('priority', 'N/A')}"
                    )
                    initiatives = gap.get("initiatives") or []
                    if initiatives:
                        for init in initiatives[:3]:
                            doc.add_paragraph(f"   • {init}")
    else:
        doc.add_paragraph("No value creation plan available.")

    # ── Save ───────────────────────────────────────────────────────────────────
    output_path = f"ic_memo_{company_ticker}.docx"
    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    asyncio.run(generate_ic_memo("NVDA"))
