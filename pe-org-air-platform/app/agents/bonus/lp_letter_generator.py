"""
LP Letter Generator (Bonus) - Generates a formal Limited Partner letter
from a completed due diligence workflow result.
"""
import asyncio
from datetime import datetime

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from exercises.agentic_due_diligence import run_due_diligence
from app.services.integration.portfolio_data_service import portfolio_data_service


def _fmt_score(value, decimals: int = 1) -> str:
    if value is None:
        return "N/A"
    try:
        return f"{float(value):.{decimals}f}"
    except (TypeError, ValueError):
        return str(value)


def _investment_thesis(org_air: float) -> str:
    if org_air >= 80: 
        return (
            f"With an Org-AI-R score of {org_air:.1f}, the company demonstrates "
            "exceptional organisational AI readiness. We have high conviction in its "
            "ability to deploy AI at scale, accelerate margin expansion, and sustain "
            "competitive differentiation over our investment horizon."
        )
    elif org_air >= 60:
        return (
            f"With an Org-AI-R score of {org_air:.1f}, the company shows solid "
            "organisational AI readiness with identifiable upside. Our value creation "
            "plan targets specific capability gaps that, if closed, are expected to "
            "drive meaningful EBITDA improvement within the hold period."
        )
    elif org_air >= 40:
        return (
            f"With an Org-AI-R score of {org_air:.1f}, the company is at an early "
            "stage of AI readiness. The investment thesis centres on structured "
            "capability building across talent, data infrastructure, and governance — "
            "with a targeted improvement to a score above 65 within 24 months."
        )
    else:
        return (
            f"With an Org-AI-R score of {org_air:.1f}, the company presents a "
            "significant transformation opportunity. Meaningful operational and "
            "structural changes are required; the investment carries elevated "
            "execution risk that is reflected in our underwriting assumptions."
        )


async def generate_lp_letter(company_id: str, assessment_type: str = "full") -> str:
    # Look up company name and ticker for display
    company_name = company_id
    company_ticker = company_id
    try:
        views = await portfolio_data_service.get_portfolio_view("growth_fund_v")
        for v in views:
            if v.company_id == company_id:
                company_name = v.name
                company_ticker = v.ticker
                break
    except Exception:
        pass

    result = await run_due_diligence(company_id, assessment_type)

    doc = Document()
    today = datetime.utcnow().strftime("%B %d, %Y")

    # ── Date ──────────────────────────────────────────────────────────────────
    date_para = doc.add_paragraph(today)
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph()

    # ── Salutation / heading ──────────────────────────────────────────────────
    doc.add_paragraph("To Our Limited Partners,")
    doc.add_paragraph()

    subject = doc.add_paragraph()
    subject_run = subject.add_run(f"Re: Portfolio Update — {company_name} ({company_ticker})")
    subject_run.bold = True

    doc.add_paragraph()

    # ── Opening ───────────────────────────────────────────────────────────────
    doc.add_paragraph(
        f"We are pleased to provide you with an update on our due diligence assessment "
        f"of {company_name}. The following letter summarises the key findings from our "
        f"Org-AI-R evaluation and our forward-looking value creation thesis."
    )
    doc.add_paragraph()

    # ── Org-AI-R score & thesis ───────────────────────────────────────────────
    doc.add_heading("Org-AI-R Assessment", level=1)

    scoring = result.get("scoring_result") or {}
    org_air_raw = scoring.get("org_air")
    try:
        org_air = float(org_air_raw) if org_air_raw is not None else 0.0
    except (TypeError, ValueError):
        org_air = 0.0

    score_table = doc.add_table(rows=4, cols=2)
    score_table.style = "Light List Accent 1"
    rows = score_table.rows
    rows[0].cells[0].text = "Org-AI-R Score"
    rows[0].cells[1].text = _fmt_score(scoring.get("org_air"))
    rows[1].cells[0].text = "Value Readiness (VR)"
    rows[1].cells[1].text = _fmt_score(scoring.get("vr_score"))
    rows[2].cells[0].text = "Human Readiness (HR)"
    rows[2].cells[1].text = _fmt_score(scoring.get("hr_score"))
    rows[3].cells[0].text = "Synergy Score"
    rows[3].cells[1].text = _fmt_score(scoring.get("synergy_score"))

    doc.add_paragraph()
    doc.add_heading("Investment Thesis", level=2)
    doc.add_paragraph(_investment_thesis(org_air))
    doc.add_paragraph()

    # ── Key findings: SEC analysis ────────────────────────────────────────────
    doc.add_heading("Key Findings", level=1)

    sec = result.get("sec_analysis") or {}
    findings = sec.get("findings") or []
    dimensions_covered = sec.get("dimensions_covered") or []

    doc.add_heading("Regulatory & Market Intelligence", level=2)
    # findings may be a dict {"count": N, "items": [...]} or a list
    evidence_items = []
    if isinstance(findings, dict):
        evidence_items = findings.get("items") or []
    elif isinstance(findings, list):
        evidence_items = findings

    if evidence_items:
        for item in evidence_items[:5]:
            if isinstance(item, dict):
                title = item.get("title") or item.get("description", "")[:80] or "Evidence signal"
                category = item.get("category", "")
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(f"[{category}] ").bold = True
                p.add_run(title[:120])
            else:
                doc.add_paragraph(str(item)[:120], style="List Bullet")
    else:
        doc.add_paragraph("No SEC findings available at this time.")

    if dimensions_covered:
        doc.add_paragraph(
            f"Dimensions assessed: {', '.join(str(d) for d in dimensions_covered)}."
        )

    doc.add_paragraph()

    # ── Key findings: evidence justifications ─────────────────────────────────
    doc.add_heading("Capability Assessment by Dimension", level=2)
    evidence = result.get("evidence_justifications") or {}
    justifications = evidence.get("justifications") or {}
    if justifications:
        for dimension, data in justifications.items():
            label = dimension.replace("_", " ").title()
            p = doc.add_paragraph(style="List Bullet")
            if isinstance(data, dict):
                score_str = _fmt_score(data.get("score"))
                strength = data.get("strength", "N/A")
                p.add_run(f"{label}: ").bold = True
                p.add_run(f"Score {score_str} | Evidence strength: {strength}")
            else:
                p.add_run(f"{label}: ").bold = True
                p.add_run(str(data))
    else:
        doc.add_paragraph("Detailed capability justifications are pending.")

    doc.add_paragraph()

    # ── Value creation highlights ─────────────────────────────────────────────
    doc.add_heading("Value Creation Highlights", level=1)
    vc_plan = result.get("value_creation_plan") or {}
    gap_data = vc_plan.get("gap_analysis") or {}
    if gap_data:
        scalar_keys = ["target_org_air", "gap_count", "top_priority", "estimated_investment", "projected_ebitda_pct"]
        for key in scalar_keys:
            if key in gap_data:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(f"{key.replace('_', ' ').title()}: ").bold = True
                p.add_run(str(gap_data[key]))
        gaps = gap_data.get("gaps") or []
        if gaps:
            doc.add_heading("Top Gaps", level=2)
            for gap in gaps[:5]:
                if isinstance(gap, dict):
                    dim = gap.get("dimension", "").replace("_", " ").title()
                    g = doc.add_paragraph(style="List Bullet")
                    g.add_run(f"{dim}: ").bold = True
                    g.add_run(
                        f"Gap {gap.get('gap', 'N/A')} pts | Priority: {gap.get('priority', 'N/A')}"
                    )
                    initiatives = gap.get("initiatives") or []
                    if initiatives:
                        doc.add_paragraph(f"  Actions: {'; '.join(initiatives[:2])}")
    else:
        doc.add_paragraph("Value creation plan is being finalised.")

    doc.add_paragraph()

    # ── HITL / governance ─────────────────────────────────────────────────────
    doc.add_heading("Investment Committee Governance", level=1)
    requires_approval = result.get("requires_approval", False)
    approval_status = result.get("approval_status") or "N/A"
    approved_by = result.get("approved_by") or "N/A"

    if requires_approval:
        doc.add_paragraph(
            f"This investment required escalated Investment Committee review. "
            f"Approval status: {approval_status}. Reviewed by: {approved_by}."
        )
    else:
        doc.add_paragraph(
            "This investment was processed under standard approval thresholds and "
            "did not require escalated Investment Committee review."
        )

    doc.add_paragraph()

    # ── Closing ───────────────────────────────────────────────────────────────
    doc.add_paragraph(
        "We remain committed to rigorous, data-driven evaluation of every portfolio "
        "company and will continue to update you as our assessment progresses. "
        "Please do not hesitate to contact your relationship manager should you have "
        "any questions regarding this update."
    )
    doc.add_paragraph()
    doc.add_paragraph("Yours sincerely,")
    doc.add_paragraph()
    doc.add_paragraph("The Investment Team")
    doc.add_paragraph("PE Org-AI-R Platform")

    # ── Save ──────────────────────────────────────────────────────────────────
    output_path = f"lp_letter_{company_ticker}.docx"
    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    asyncio.run(generate_lp_letter("NVDA"))
