
from docx import Document
import os

def update_docx(file_path):
    if not os.path.exists(file_path):
        print(f"File not found: {file_path}")
        return

    doc = Document(file_path)
    
    # Update Links
    for p in doc.paragraphs:
        if "https://codelabs-preview.appspot.com/" in p.text:
            p.text = p.text.replace("#0", "#6")

    # 1. Update Introduction
    for p in doc.paragraphs:
        if "SEC filings, job postings, patents, and technology stacks" in p.text:
            p.text = p.text.replace(
                "SEC filings, job postings, patents, and technology stacks",
                "SEC filings, job postings, patents, technology stacks, board composition (Case Study 3), and Glassdoor employee reviews (Case Study 3)"
            )
        
        # 2. Update Signal Collection steps
        if "Technology stack detection" in p.text and "backfill" in p.text.lower():
            p.text = p.text + "\n☐ Board composition and leadership diversity audit\n☐ Glassdoor cultural sentiment and AI awareness sweep"

        # 3. Update Architecture steps
        if "4️ Signal Collectors → Parallel execution (jobs, patents, tech stack)" in p.text:
            p.text = p.text.replace(
                "(jobs, patents, tech stack)",
                "(jobs, patents, tech stack, board composition, and culture)"
            )

        # 4. Update Scoring Logic details
        if "Aggregates signals into maturity dimensions" in p.text:
            p.text = p.text.replace(
                "Aggregates signals into maturity dimensions",
                "Aggregates signals into maturity dimensions using calibrated industry bases and dynamic Position Factor (PF) logic"
            )

    # 5. Robust Airflow Section
    # Find a good place to insert Airflow details. Maybe after "Understanding the Architecture"
    found_arch = False
    for i, p in enumerate(list(doc.paragraphs)):
        if "Understanding the Architecture" in p.text:
            found_arch = True
            # Insert after the next few paragraphs or at the end of section
            
    # Add a new section for Airflow if it's missing or thin
    airflow_content = [
        "🔄 Airflow Orchestration & Pipeline Resilience",
        "The platform utilizes Apache Airflow 2.x to manage complex data dependencies and ensure horizontal scalability.",
        "Key Orchestration Features:",
        "• Dynamic Task Mapping: The pipeline automatically scales to N companies using Airflow's .expand() pattern, enabling parallel processing of the entire portfolio.",
        "• Asynchronous Bridge: The FastAPI backend acts as a singleton gateway, triggering DAGs via the Airflow REST API while maintaining real-time status tracking in the dashboard.",
        "• Trigger Rules & Error Handling: Mapped tasks utilize TriggerRule.ALL_DONE, ensuring that a partial failure in one signal (e.g., a scraper block) doesn't stop the overall calculation for a company.",
        "• XCom Backend Optimization: High-volume data (like SEC chunks) are passed via shared volumes or S3, keeping the Airflow metadata database lean and performant."
    ]
    
    # Just append to the end for now or find the Airflow section if it exists
    doc.add_heading("8. Airflow Pipeline Orchestration", level=2)
    for line in airflow_content:
        doc.add_paragraph(line)

    # Better approach for table updates
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "patents, and tech stack" in cell.text:
                    cell.text = cell.text.replace("patents, and tech stack", "patents, tech stack, board, and culture")

    output_path = file_path.replace(".docx", "_Updated_v2.docx")
    doc.save(output_path)
    print(f"Updated document saved to: {output_path}")

if __name__ == "__main__":
    update_docx("../PE Org-AI-R Platform_ Intelligence Engine for Due Diligence.docx")
